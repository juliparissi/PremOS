from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "PremOS_manual_implementacion_comercial.docx"
LOGO = ROOT / "public" / "icon.png"


BLUE = "1F4D78"
ACCENT = "06B6D4"
INK = "111827"
MUTED = "64748B"
LIGHT = "E8EEF5"
SOFT = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    r.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths_cm):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def add_table(doc, headers, rows, widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True, color="0B2545")
        set_cell_shading(hdr[idx], LIGHT)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))

    if widths_cm:
        set_table_widths(table, widths_cm)
    doc.add_paragraph()
    return table


def add_callout(doc, title, body, fill=SOFT):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph()


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "0B1220")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(text.strip().splitlines()):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string("E5E7EB")
    doc.add_paragraph()


def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_step(doc, number, title, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{number}. {title}")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    r.font.size = Pt(11)
    p2 = doc.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(4)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("PremOS - Manual de implementacion comercial")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def build():
    doc = Document()
    setup_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    add_footer(section)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        title.add_run().add_picture(str(LOGO), width=Inches(0.75))
    title.add_run().add_break()
    r = title.add_run("PremOS")
    r.bold = True
    r.font.size = Pt(27)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    title.add_run().add_break()
    sub = title.add_run("Manual de implementacion comercial")
    sub.font.size = Pt(14)
    sub.font.color.rgb = RGBColor.from_string(MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Version operativa - Junio 2026")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    add_callout(
        doc,
        "Objetivo del manual",
        "Estandarizar la instalacion de PremOS para nuevos clientes: preparar Supabase, configurar el VPS, activar licencia, dejar modulos correctos por plan y revisar rapido los problemas mas comunes.",
        fill="EEF6FF",
    )

    doc.add_heading("1. Flujo comercial recomendado", level=1)
    add_step(doc, 1, "Relevar el negocio", "Registrar rubro, productos, necesidades, dominio deseado, logo, datos fiscales y plan contratado.")
    add_step(doc, 2, "Crear entorno tecnico", "Preparar un proyecto Supabase por cliente y una carpeta/app separada en el VPS.")
    add_step(doc, 3, "Activar licencia", "Cargar el cliente en el panel administrador, asignar plan, estado, vencimiento y license_key.")
    add_step(doc, 4, "Configurar PremOS", "Cargar variables de entorno, compilar, iniciar con PM2 y conectar dominio/SSL.")
    add_step(doc, 5, "Validar y entregar", "Probar flujo completo y entregar URL, usuario temporal, plan activo y canal de soporte.")

    doc.add_heading("2. Datos que pedir al cliente", level=1)
    add_checklist(doc, [
        "Nombre comercial, razon social, CUIT, direccion, telefono y email.",
        "Logo en PNG o JPG.",
        "Lista inicial de productos, modelos, colores y precios de referencia.",
        "Suministros principales y unidades de medida.",
        "Proceso de produccion y recetas si contrata Full o Pro.",
        "Dominio o subdominio deseado, por ejemplo app.empresa.com.",
        "Plan contratado: Lite, Full o Pro.",
        "Necesita IA: si, no, o etapa futura.",
    ])

    doc.add_heading("3. Planes y modulos", level=1)
    add_table(
        doc,
        ["Plan", "Modulos incluidos", "Uso recomendado"],
        [
            ["Lite", "Resumen, Clientes, Productos, Presupuestos, Ventas, Pedidos, Economia, Soporte, Configuracion", "Primera version para comercios que necesitan ordenar ventas, presupuestos y caja."],
            ["Full", "Todo Lite + Reportes, Lista de precios, Produccion, Nomina, Suministro y Stock", "Fabricas que necesitan controlar materias primas, personal, produccion, inventario y reportes."],
            ["Pro", "Todo Full + Asistente IA", "Clientes que quieren consultas asistidas y lectura rapida de indicadores del negocio."],
        ],
        [2.0, 8.3, 5.8],
    )
    add_callout(doc, "Regla comercial", "El plan real debe manejarse desde el panel administrador. En instalaciones con licencia, PremOS toma el plan desde el panel y no desde Configuracion local.")

    doc.add_heading("4. Crear Supabase del cliente", level=1)
    add_checklist(doc, [
        "Crear un proyecto nuevo de Supabase para el cliente.",
        "Guardar Project URL, anon/publishable key y service_role key.",
        "Ir a SQL Editor.",
        "Ejecutar el archivo supabase/premos_schema_instalacion.sql.",
        "Crear el usuario inicial en Authentication.",
        "Probar que el login ingrese a PremOS.",
    ])
    add_table(
        doc,
        ["Tabla principal", "Para que se usa"],
        [
            ["clientes, productos, colores", "Base comercial para presupuestos, ventas y pedidos."],
            ["presupuestos, presupuesto_items", "Cotizaciones y detalle de productos cotizados."],
            ["pedidos, pedido_items, pagos_pedidos", "Seguimiento de pedidos y cobranzas."],
            ["movimientos_economia, movimientos_economia_abonos, proveedores", "Caja, salidas de dinero, pagos a proveedores y gastos."],
            ["nomina_empleados, nomina_movimientos", "Empleados, sueldos, adelantos e impacto automatico en Economia."],
            ["stock, suministros, movimientos_suministro", "Inventario comercial y materias primas."],
            ["produccion, produccion_items, recetas_produccion", "Ordenes de produccion, consumo de insumos y trazabilidad."],
            ["listas_precios, lista_precios_items", "Listas comerciales exportables a PDF o Excel."],
            ["configuracion_empresa, notas_rapidas", "Datos de empresa, PDF y notas del resumen."],
        ],
        [5.0, 11.1],
    )

    doc.add_heading("5. Variables de entorno de PremOS", level=1)
    add_code(doc, """
NEXT_PUBLIC_SUPABASE_URL=https://xxxx.supabase.co
NEXT_PUBLIC_SUPABASE_ANON_KEY=sb_publishable_xxxxx
SUPABASE_SERVICE_ROLE_KEY=xxxxx

OPENAI_API_KEY=

NEXT_PUBLIC_PREMOS_LICENSE_API_URL=https://admin.premoos.site/api/licencias
NEXT_PUBLIC_PREMOS_LICENSE_KEY=cliente_license_key
NEXT_PUBLIC_PREMOS_SUPPORT_API_URL=https://admin.premoos.site/api/soporte/tickets
""")
    add_callout(
        doc,
        "Importante",
        "No subir nunca .env.local a GitHub. La anon/publishable key va en el frontend; la service_role key y OPENAI_API_KEY deben tratarse como privadas.",
        fill="FFF7ED",
    )

    doc.add_heading("6. Instalacion desde Git en VPS", level=1)
    add_code(doc, """
ssh -p5507 root@IP_DEL_VPS
cd /root
git clone git@github.com:juliparissi/PremOS.git cliente-premos
cd cliente-premos
npm install
npm run build
pm2 start npm --name cliente-premos -- start -- -H 0.0.0.0 -p 3002
pm2 save
""")
    add_callout(doc, "Puertos", "Usar un puerto distinto por entorno. Ejemplo: PremOS propio en 3000, demo en 3001, cliente nuevo en 3002.")

    doc.add_heading("7. Actualizar un cliente existente", level=1)
    add_code(doc, """
ssh -p5507 root@IP_DEL_VPS
cd /root/cliente-premos
git pull origin main
npm install
npm run build
pm2 restart cliente-premos
pm2 logs cliente-premos --lines 30
""")

    doc.add_heading("8. Nginx y SSL", level=1)
    add_code(doc, """
sudo nano /etc/nginx/sites-available/cliente-premos

server {
    server_name app.empresa.com;

    location / {
        proxy_pass http://127.0.0.1:3002;
        proxy_http_version 1.1;
        proxy_set_header Upgrade $http_upgrade;
        proxy_set_header Connection 'upgrade';
        proxy_set_header Host $host;
        proxy_cache_bypass $http_upgrade;
    }
}

sudo ln -s /etc/nginx/sites-available/cliente-premos /etc/nginx/sites-enabled/
sudo nginx -t
sudo systemctl reload nginx
sudo certbot --nginx -d app.empresa.com
""")

    doc.add_heading("9. Panel administrador y licencias", level=1)
    add_checklist(doc, [
        "Crear cliente en el panel admin.",
        "Asignar plan: Lite, Full o Pro.",
        "Asignar estado: Activo, Demo, Suspendido o Vencido.",
        "Definir vencimiento mensual.",
        "Copiar license_key al .env.local de PremOS del cliente.",
        "Verificar que el sistema muestre solo los modulos del plan.",
        "Si el cliente no paga, cambiar estado a Suspendido o dejar vencer la licencia.",
    ])
    add_table(
        doc,
        ["Estado", "Efecto esperado"],
        [
            ["Activo", "El cliente puede usar el sistema normalmente."],
            ["Demo", "Permite mostrar el sistema por tiempo limitado."],
            ["Suspendido", "PremOS bloquea el acceso y muestra servicio suspendido."],
            ["Vencido", "PremOS bloquea el acceso cuando pasa la fecha de vencimiento."],
        ],
        [3.0, 13.1],
    )

    doc.add_heading("10. Soporte por tickets", level=1)
    add_checklist(doc, [
        "El cliente crea tickets desde el modulo Soporte.",
        "El panel admin recibe tickets asociados al cliente por license_key.",
        "Responder desde el panel admin.",
        "El cliente puede continuar la conversacion hasta que el ticket se cierre.",
        "Los tickets cerrados deben limpiarse automaticamente luego de 72 horas segun la regla del panel.",
    ])

    doc.add_heading("11. Configuracion inicial dentro del sistema", level=1)
    add_checklist(doc, [
        "Entrar con el usuario del cliente.",
        "Cargar datos de empresa en Configuracion.",
        "Subir logo de empresa para PDFs.",
        "Crear clientes, productos y colores iniciales.",
        "Cargar listas de precios si el plan lo permite.",
        "Cargar proveedores si usara Economia.",
        "Cargar empleados si usara Nomina.",
        "Cargar suministros y recetas si el plan incluye Produccion.",
        "Probar PDF de presupuesto, nota de venta y reportes.",
    ])

    doc.add_heading("12. Prueba final antes de entregar", level=1)
    add_table(
        doc,
        ["Prueba", "Resultado esperado"],
        [
            ["Login", "El usuario accede y el plan muestra los modulos correctos."],
            ["Presupuesto", "Se crea, descarga PDF y puede pasar a pedido una sola vez."],
            ["Venta directa", "Crea pedido, no habilita PDF de presupuesto y registra seña si corresponde."],
            ["Pedido", "Permite registrar cobros desde el modal del pedido."],
            ["Economia", "Muestra ingresos desde ventas/pedidos y permite solo salidas de caja manuales."],
            ["Nomina", "Permite cargar empleados y cada pago genera una salida en Economia."],
            ["Produccion", "Descuenta suministros y permite eliminar produccion devolviendo insumos."],
            ["Reportes", "Muestra importes correctos y permite ocultar saldos."],
            ["Soporte", "Ticket enviado desde PremOS aparece en panel admin."],
        ],
        [4.1, 12.0],
    )

    doc.add_heading("13. Revisiones rapidas si algo falla", level=1)
    add_table(
        doc,
        ["Problema", "Revision rapida"],
        [
            ["No abre el dominio", "Revisar DNS, firewall del cloud, Nginx, SSL y puerto PM2."],
            ["PM2 aparece en errored", "Ejecutar pm2 logs NOMBRE --lines 50 y revisar puerto ocupado o falta de .env."],
            ["Error EADDRINUSE", "Ese puerto ya esta en uso. Cambiar puerto o detener proceso duplicado."],
            ["No conecta Supabase", "Revisar NEXT_PUBLIC_SUPABASE_URL y NEXT_PUBLIC_SUPABASE_ANON_KEY."],
            ["No funciona IA", "Revisar OPENAI_API_KEY y que el plan sea Pro si se vende como modulo activo."],
            ["No cambian modulos por plan", "Revisar NEXT_PUBLIC_PREMOS_LICENSE_API_URL y NEXT_PUBLIC_PREMOS_LICENSE_KEY."],
            ["No llegan tickets", "Revisar NEXT_PUBLIC_PREMOS_SUPPORT_API_URL y license_key del cliente."],
            ["PDF sale con datos viejos", "Revisar Configuracion de empresa y limpiar cache del navegador si corresponde."],
        ],
        [4.4, 11.7],
    )

    doc.add_heading("14. Comandos utiles", level=1)
    add_code(doc, """
pm2 list
pm2 logs cliente-premos --lines 50
pm2 restart cliente-premos
pm2 stop cliente-premos
pm2 delete cliente-premos
pm2 save

sudo nginx -t
sudo systemctl reload nginx

git status
git pull origin main
npm run build
""")

    doc.add_heading("15. Checklist de entrega", level=1)
    add_checklist(doc, [
        "URL final funcionando con HTTPS.",
        "Usuario y contrasena temporal entregados.",
        "Datos de empresa cargados.",
        "Plan y vencimiento confirmados en panel admin.",
        "Licencia activa probada.",
        "Soporte por tickets probado.",
        "Backups y renovacion mensual registrados.",
        "Cliente informado sobre alcance del plan contratado.",
    ])

    doc.add_heading("16. Politica operativa recomendada", level=1)
    doc.add_paragraph(
        "Para la primera etapa comercial, conviene mantener una instalacion separada por cliente: un proyecto Supabase por cliente y una app PM2 independiente. "
        "Esto simplifica privacidad, soporte, bloqueo por falta de pago y migraciones. Mas adelante, cuando el volumen lo justifique, se puede evolucionar hacia un esquema multi-tenant o infraestructura dedicada."
    )

    doc.add_heading("17. Habilitar u ocultar modulos", level=1)
    doc.add_paragraph(
        "Los modulos visibles por plan se controlan desde el archivo lib/planes.ts. "
        "Si el panel administrador esta activo, el cliente recibe su plan desde la licencia; PremOS toma ese plan y muestra solo los modulos permitidos."
    )
    add_table(
        doc,
        ["Archivo / bloque", "Que tocar"],
        [
            ["PremosModule", "Agregar el identificador interno del modulo, por ejemplo nomina."],
            ["planModules", "Agregar o quitar el modulo dentro de lite, full o pro."],
            ["routeModules", "Vincular la ruta del modulo con su identificador, por ejemplo /nomina -> nomina."],
            ["app/layout.tsx", "Solo revisar si el modulo no aparece en la barra lateral o si necesita etiqueta especial."],
            ["Panel admin", "Cambiar el plan del cliente. No hace falta tocar PremOS para pasar de Lite a Full o Pro."],
        ],
        [4.0, 12.1],
    )
    add_code(doc, """
# 1. Editar el archivo de planes
lib/planes.ts

# 2. Validar que compile
npm run build

# 3. Subir a GitHub
git status
git add lib/planes.ts
git commit -m "chore: ajustar modulos por plan"
git push origin main

# 4. Actualizar el VPS del cliente
ssh -p5507 root@IP_DEL_VPS
cd /root/cliente-premos
git pull origin main
npm install
npm run build
pm2 restart cliente-premos
""")
    add_callout(
        doc,
        "Regla simple",
        "Para ocultar un modulo a un plan, quitarlo del array correspondiente en planModules. Para habilitarlo, agregarlo. Si el modulo tiene ruta propia, tambien debe existir en routeModules.",
        fill="EEF6FF",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
